import re
from html.parser import HTMLParser
from io import BytesIO
from typing import List, Optional, Tuple

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

WHITESPACE_RE = re.compile(r"\s+")


class HTML2Docx(HTMLParser):
    def __init__(self, title: str):
        super().__init__()
        self.doc = Document()
        self.doc.core_properties.title = title
        self.list_style: List[str] = []
        self.href = ""
        self._reset()

    def _reset(self) -> None:
        self.p: Optional[Paragraph] = None
        self.r: Optional[Run] = None

        # Formatting options
        self.pre = False
        self.attrs: List[List[str]] = []
        self.collapse_space = True

    def finish_p(self) -> None:
        if self.r is not None:
            self.r.text = self.r.text.rstrip()
        self._reset()

    def init_run(self, attrs: List[str]) -> None:
        self.r = None
        self.attrs.append(attrs)

    def finish_run(self) -> None:
        self.attrs = self.attrs[:-1]
        self.r = None

    def add_text(self, data: str) -> None:
        if self.p is None:
            style = self.list_style[-1] if self.list_style else None
            self.p = self.doc.add_paragraph(style=style)
        if self.r is None:
            self.r = self.p.add_run()
            for attrs in self.attrs:
                for run_style in attrs:
                    setattr(self.r.font, run_style, True)
        self.r.add_text(data)

    def add_list_style(self, name: str) -> None:
        self.finish_p()
        # The template included by python-docx only has 3 list styles.
        level = min(len(self.list_style) + 1, 3)
        suffix = f" {level}" if level > 1 else ""
        self.list_style.append(f"{name}{suffix}")

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        if tag == "a":
            self.href = str(next((val for name, val in attrs if name == "href"), ""))
            self.init_run([])
        elif tag in ["b", "strong"]:
            self.init_run(["bold"])
        elif tag == "br":
            if self.r:
                self.r.add_break()
        elif tag in ["em", "i"]:
            self.init_run(["italic"])
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag[-1])
            self.p = self.doc.add_heading(level=level)
        elif tag == "ol":
            self.add_list_style("List Number")
        elif tag == "pre":
            self.pre = True
        elif tag == "u":
            self.init_run(["underline"])
        elif tag == "ul":
            self.add_list_style("List Bullet")

    def handle_data(self, data: str) -> None:
        if not self.pre:
            data = re.sub(WHITESPACE_RE, " ", data)
        if self.collapse_space:
            data = data.lstrip()
        if data:
            if self.href:
                if data.endswith(" "):
                    data += self.href + " "
                else:
                    data += " " + self.href
                self.href = ""
            self.collapse_space = data.endswith(" ")
            self.add_text(data)

    def handle_endtag(self, tag: str) -> None:
        if tag in ["a", "b", "em", "i", "strong", "u"]:
            self.finish_run()
        elif tag in ["h1", "h2", "h3", "h4", "h5", "h6", "li", "ol", "p", "pre", "ul"]:
            self.finish_p()
            if tag in ["ol", "ul"]:
                del self.list_style[-1]
            elif tag == "pre":
                self.pre = False


def html2docx(content: str, title: str) -> BytesIO:
    """Convert valid HTML content to a docx document and return it as a
    io.BytesIO() object.
    """
    parser = HTML2Docx(title)
    parser.feed(content.strip())

    buf = BytesIO()
    parser.doc.save(buf)
    return buf
