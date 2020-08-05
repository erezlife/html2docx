import os

import docx

from html2docx import html2docx

from .utils import TEST_DIR


def strip_ns(x):
    pos = x.find("}") + 1
    return x[pos:]


def attrib(d, key):
    for k, v in d.items():
        if k.endswith("}" + key):
            return v


def test_links():
    html_path = os.path.join(TEST_DIR, "links.html")
    html = open(html_path).read()
    buf = html2docx(html, title="links")

    doc = docx.Document(buf)

    assert len(doc.paragraphs) == 9

    # check hyperlink
    run = doc.paragraphs[0]._p[1]
    assert strip_ns(run.tag) == "hyperlink"
    assert attrib(run.attrib, "anchor") == "1.1"

    children = run.getchildren()
    assert len(children) == 1

    wR = children[0]
    assert strip_ns(wR.tag) == "r"
    assert wR.text == "link to bookmark"

    # check bookmark
    run = doc.paragraphs[6]._p[1]
    assert strip_ns(run.tag) == "bookmarkStart"

    run = doc.paragraphs[6]._p[2]
    assert strip_ns(run.tag) == "r"
    assert run.text == "1.1 Bookmark"

    run = doc.paragraphs[6]._p[3]
    assert strip_ns(run.tag) == "bookmarkEnd"
