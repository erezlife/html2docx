import os

import docx

from html2docx import html2docx

from .utils import TEST_DIR


def test_table():
    html_path = os.path.join(TEST_DIR, "table.html")
    html = open(html_path).read()
    buf = html2docx(html, title="table")

    doc = docx.Document(buf)

    assert len(doc.tables) == 1
    table = doc.tables[0]

    assert len(table.rows) == 2
    assert len(table.columns) == 2

    contents = [
        ["1", "2"],
        ["3"],
    ]

    for r, row in enumerate(contents):
        for c, text in enumerate(row):
            assert table.cell(r, c).text == text

    assert table.cell(0, 0).paragraphs[0].runs[0].font.bold is None
    assert table.cell(0, 1).paragraphs[0].runs[0].font.bold is True


def test_table_nested():
    html_path = os.path.join(TEST_DIR, "table_nested.html")
    html = open(html_path).read()
    buf = html2docx(html, title="table")

    doc = docx.Document(buf)
    assert len(doc.tables) == 1

    table = doc.tables[0]

    assert len(table.rows) == 2
    assert len(table.columns) == 2

    contents = [
        ["1", "2"],
        ["3"],
    ]

    for r, row in enumerate(contents):
        for c, text in enumerate(row):
            assert table.cell(r, c).text == text

    cell = table.cell(1, 1)
    assert len(cell.tables) == 1

    table2 = cell.tables[0]

    contents2 = [
        ["4", "5"],
        ["6", "7"],
    ]

    for r, row in enumerate(contents2):
        for c, text in enumerate(row):
            assert table2.cell(r, c).text == text
