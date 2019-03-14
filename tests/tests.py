import json
import pathlib
from unittest import TestCase

import docx

from html2docx import html2docx


class Html2DocxTest(TestCase):
    def test_html2docx(self):
        datadir = pathlib.Path(__file__).parent.joinpath("html2docx")
        for path in datadir.glob("*.html"):
            with self.subTest(path.name):
                with datadir.joinpath(f"{path.stem}.json").open() as fp:
                    spec = json.load(fp)

                with path.open() as fp:
                    html = fp.read()

                buf = html2docx(html, title=path.name)
                doc = docx.Document(buf)

                self.assertEqual(doc.core_properties.title, path.name)
                self.assertEqual(len(doc.paragraphs), len(spec))
                for p, p_spec in zip(doc.paragraphs, spec):
                    self.assertEqual(p.text, p_spec["text"])
                    self.assertEqual(p.style.name, p_spec.get("style", "Normal"))

                    runs_spec = p_spec["runs"]
                    self.assertEqual(len(p.runs), len(runs_spec))
                    for run, run_spec in zip(p.runs, runs_spec):
                        self.assertEqual(run.text, run_spec["text"])
                        self.assertEqual(
                            run.bold,
                            run_spec.get("bold", False),
                            f"Wrong bold for text '{run.text}'.",
                        )
                        self.assertEqual(
                            run.italic,
                            run_spec.get("italic", False),
                            f"Wrong italic for text '{run.text}'.",
                        )
                        self.assertEqual(
                            run.underline,
                            run_spec.get("underline", False),
                            f"Wrong underline for text '{run.text}'.",
                        )
